import os
import re
import json
import asyncio
from typing import List, Optional
import requests
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

app = FastAPI(title="JurisPrime API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
CNJ_API_KEY = os.getenv("CNJ_API_KEY", "")

# --- MÓDULO DATAJUD / CNJ ---
def consultar_datajud(numero_processo: str, tribunal: str = "tjsp") -> Optional[str]:
    if not CNJ_API_KEY:
        return None
    num_limpo = re.sub(r"\D", "", numero_processo)
    if len(num_limpo) != 20:
        return None
    
    url = f"https://api-publica.datajud.cnj.jus.br/api_publica_{tribunal.lower()}/_search"
    headers = {
        "Authorization": f"APIKey {CNJ_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {"query": {"match": {"numeroProcesso": num_limpo}}}
    
    try:
        res = requests.post(url, json=payload, headers=headers, timeout=8)
        if res.status_code == 200:
            hits = res.json().get("hits", {}).get("hits", [])
            if hits:
                proc = hits[0].get("_source", {})
                classe = proc.get("classe", {}).get("nome", "Não informada")
                orgao = proc.get("orgaoJulgador", {}).get("nome", "Não informado")
                assuntos = [a.get("nome", "") for a in proc.get("assuntos", [])]
                return f"[DADOS OFICIAIS CNJ/{tribunal.upper()}]: Classe: {classe} | Vara/Órgão: {orgao} | Assuntos: {', '.join(assuntos)}"
    except Exception:
        return None
    return None

# --- PROMPT FORENSE DE 1º GRAU ---
SUPERPROMPT_PETICAO_1GRAU = """
Você é um Advogado Sênior e Especialista em Direito Processual Civil e Prática Forense de 1º Grau.
Sua missão é redigir uma PETIÇÃO INICIAL DE 1º GRAU (ou Peça Processual Técnica) completa, exaustiva, de alta densidade jurídica e pronta para protocolo (meta de 2.000 a 3.500 palavras).

DIRETRIZES TÉCNICAS E FORENSES:
1. ENDEREÇAMENTO PRECISO: Ao d. Juízo da Vara Cível / Juizado Especial da Comarca competente.
2. QUALIFICAÇÃO DAS PARTES: Formato forense completo com indicação de requerimento de benefícios (Justiça Gratuita, Prioridade de Tramitação, se aplicável).
3. FATOS CRONOLÓGICOS E PORMENORIZADOS: Narrativa estruturada e clara, indicando a relação jurídica, conduta lesiva, dano e nexo causal.
4. TUTELA DE URGÊNCIA / EVIDÊNCIA (Art. 300 / 311 do CPC): Se solicitada ou aplicável, fundamente exaustivamente a probabilidade do direito (fumus boni iuris) e o perigo de dano (periculum in mora), com pedido liminar expresso inaudita altera parte.
5. FUNDAMENTAÇÃO JURÍDICA ROBUSTA: Articulação do Código Civil, CPC, CDC e precedentes consolidados (Súmulas e Temas Repetitivos do STJ/STF).
6. PEDIDOS E REQUERIMENTOS FINAIS: Relação minuciosa com citações, produção de provas, inversão do ônus da prova, procedência integral, condenação em custas/sucumbência e valor da causa.
"""

@app.post("/api/peticao/gerar-stream")
async def gerar_peticao_stream(
    instrucao_usuario: str = Form(...),
    tribunal: str = Form("tjsp"),
    arquivos: List[UploadFile] = File(None)
):
    if not GEMINI_API_KEY:
        raise HTTPException(status_code=500, detail="Chave GEMINI_API_KEY não configurada no servidor.")
    
    client = genai.Client(api_key=GEMINI_API_KEY)
    user_parts = []
    
    # Varredura CNJ caso haja número no texto
    match_cnj = re.search(r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}", instrucao_usuario)
    if match_cnj:
        dados_cnj = consultar_datajud(match_cnj.group(0), tribunal=tribunal)
        if dados_cnj:
            user_parts.append(types.Part.from_text(text=dados_cnj))
    
    # Processamento direto de PDFs em bytes
    if arquivos:
        for file in arquivos:
            conteudo = await file.read()
            if file.content_type == "application/pdf" or file.filename.endswith(".pdf"):
                user_parts.append(types.Part.from_bytes(data=conteudo, mime_type="application/pdf"))
                user_parts.append(types.Part.from_text(text=f"[Documento Anexo: {file.filename}]"))
    
    user_parts.append(types.Part.from_text(text=instrucao_usuario))

    async def stream_generator():
        try:
            config = types.GenerateContentConfig(
                system_instruction=SUPERPROMPT_PETICAO_1GRAU,
                temperature=0.1,
                max_output_tokens=8192,
                tools=[types.Tool(google_search=types.GoogleSearch())]
            )
            response = client.models.generate_content_stream(
                model="gemini-2.5-flash",
                contents=[types.Content(role="user", parts=user_parts)],
                config=config
            )
            for chunk in response:
                if chunk.text:
                    yield f"data: {json.dumps({'text': chunk.text})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(stream_generator(), media_type="text/event-stream")

# --- EXPORTADOR DOCX PROFISSIONAL COM MARGENS FORENSES ---
class ExportDocxRequest(BaseModel):
    titulo: str
    conteudo_markdown: str

@app.post("/api/exportar-docx")
async def exportar_docx(req: ExportDocxRequest):
    doc = Document()
    
    # Margens padrão ABNT / Forense
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.18)     # 3 cm
        section.left_margin = Inches(1.18)    # 3 cm
        section.right_margin = Inches(0.78)   # 2 cm
        section.bottom_margin = Inches(0.78)  # 2 cm

    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(17, 24, 39)

    linhas = req.conteudo_markdown.split("\n")
    for linha in linhas:
        texto = linha.strip()
        if not texto:
            doc.add_paragraph()
            continue
        
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        
        if texto.startswith("# "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(texto.replace("# ", ""))
            run.bold = True
            run.font.size = Pt(14)
        elif texto.startswith("## ") or texto.startswith("### "):
            run = p.add_run(texto.replace("## ", "").replace("### ", ""))
            run.bold = True
            run.font.size = Pt(12)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.78) # Recuo forense
            p.add_run(texto)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={req.titulo}.docx"}
    )
